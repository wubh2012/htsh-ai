"""文档解析服务"""
import re
from pathlib import Path
from typing import Optional
import docx
import fitz  # pymupdf


class DocumentParser:
    """文档解析器，支持Word和PDF"""

    @staticmethod
    async def parse(file_path: str, file_type: str) -> str:
        """
        解析文档并返回纯文本
        :param file_path: 文件路径
        :param file_type: "docx" 或 "pdf"
        :return: 提取的纯文本
        :raises ValueError: 不支持的文件类型或解析失败
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if file_type.lower() == "docx":
            return await DocumentParser._parse_docx(file_path)
        elif file_type.lower() == "pdf":
            return await DocumentParser._parse_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}，仅支持 docx 和 pdf")

    @staticmethod
    async def _parse_docx(file_path: Path) -> str:
        """解析Word文档"""
        try:
            doc = docx.Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            text = "\n".join(paragraphs)
            return DocumentParser._clean_text(text)

        except Exception as e:
            raise ValueError(f"解析Word文档失败: {str(e)}")

    @staticmethod
    async def _parse_pdf(file_path: Path) -> str:
        """解析PDF文档（文字型）"""
        try:
            doc = fitz.open(file_path)
            paragraphs = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()

                # 按行分割，过滤空行
                lines = [line.strip() for line in text.split("\n") if line.strip()]
                if lines:
                    # 将同一页的内容合并，以换行分隔
                    paragraphs.extend(lines)

            doc.close()

            if not paragraphs:
                raise ValueError("PDF文档可能是扫描版或无文字内容，MVP阶段不支持")

            text = "\n".join(paragraphs)
            return DocumentParser._clean_text(text)

        except Exception as e:
            if "scan" in str(e).lower() or "image" in str(e).lower():
                raise ValueError("PDF文档可能是扫描版，MVP阶段不支持")
            raise ValueError(f"解析PDF文档失败: {str(e)}")

    @staticmethod
    def _clean_text(text: str) -> str:
        """清洗文本"""
        # 去除多余空白字符
        text = re.sub(r'[ \t]+', ' ', text)  # 多个空格合并为一个
        text = re.sub(r'\n{3,}', '\n\n', text)  # 超过2个连续换行缩减为2个
        text = re.sub(r' +\n', '\n', text)  # 行尾空格去除
        text = re.sub(r'\n +', '\n', text)  # 行首空格去除

        # 处理特殊编码字符
        text = text.replace('\u200b', '')  # 零宽空格
        text = text.replace('\ufeff', '')  # BOM字符

        # 去除不可见字符
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')

        return text.strip()

    @staticmethod
    def validate_file_type(filename: str) -> Optional[str]:
        """
        验证文件类型
        :param filename: 文件名
        :return: 文件类型 (docx/pdf) 或 None
        """
        suffix = Path(filename).suffix.lower()
        if suffix == ".docx":
            return "docx"
        elif suffix == ".pdf":
            return "pdf"
        return None
